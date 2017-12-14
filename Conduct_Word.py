import docx
file = docx.Document("D:\我的Python练习文件\MyWord.docx")
print("段落数："+str(len(file.paragraphs)))

for para in file.paragraphs:
	print(para.text)

